import io
from pypdf import PdfReader
import docx

def parse_pdf(file_bytes) -> str:
    """Extract text from PDF file bytes."""
    text_content = []
    try:
        pdf_file = io.BytesIO(file_bytes)
        reader = PdfReader(pdf_file)
        for page_num in range(len(reader.pages)):
            page = reader.pages[page_num]
            page_text = page.extract_text()
            if page_text:
                text_content.append(page_text)
        return "\n".join(text_content).strip()
    except Exception as e:
        raise ValueError(f"Error parsing PDF file: {str(e)}")

def parse_docx(file_bytes) -> str:
    """Extract text from DOCX file bytes."""
    try:
        docx_file = io.BytesIO(file_bytes)
        doc = docx.Document(docx_file)
        text_content = []
        for paragraph in doc.paragraphs:
            if paragraph.text:
                text_content.append(paragraph.text)
        # Also parse table contents
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        text_content.append(cell.text)
        return "\n".join(text_content).strip()
    except Exception as e:
        raise ValueError(f"Error parsing DOCX file: {str(e)}")

def extract_text(file_bytes, filename: str) -> str:
    """Route file bytes to the appropriate parser based on file extension."""
    ext = filename.lower().split(".")[-1]
    if ext == "pdf":
        return parse_pdf(file_bytes)
    elif ext in ["docx", "doc"]:
        return parse_docx(file_bytes)
    else:
        raise ValueError("Unsupported file format. Please upload a PDF or DOCX document.")
